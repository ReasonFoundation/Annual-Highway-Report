from docx import Document
import pandas as pd

file_path = "AHR_data 28th.xlsx"
all_sheets = pd.read_excel(file_path, sheet_name=None)

state_data = {}
for sheet_name, df in all_sheets.items():
    df = df[df['state'] != 'Total US'].copy()
    for _, row in df.iterrows():
        state = row['state']
        if state not in state_data:
            state_data[state] = {}
        state_data[state][sheet_name] = row.to_dict()

def generate_report(state, data):
    doc = Document()
    doc.add_heading(f"{state} Ranks {data['Overall']['Rank']} in the Nation in Highway Performance and Cost-Effectiveness", level=1)

    # Summary paragraph
    doc.add_paragraph(
        f"{state}'s highway system ranks {data['Overall']['Rank']} in the nation in overall cost-effectiveness and condition. "
        f"According to the Annual Highway Report, this is a change from {data['Overall']['Last Year']} last year and {data['Overall']['Five Years Ago']} five years ago."
    )

    # Pavement & safety
    doc.add_heading("Safety and Condition", level=2)
    doc.add_paragraph(
        f"{state}'s highways rank {data['Urban Interstate']['Rank']} in urban Interstate pavement condition, "
        f"{data['Rural Interstate']['Rank']} in rural Interstate pavement condition, "
        f"{data['Urban Arterial']['Rank']} in urban arterial condition, etc."
    )

    # Cost-effectiveness
    doc.add_heading("Spending and Cost-Effectiveness", level=2)
    doc.add_paragraph(
        f"{state} ranks {data['Capital & Bridge']['Rank']} in capital and bridge disbursements, "
        f"{data['Maintenance']['Rank']} in maintenance, "
        f"and {data['Administrative']['Rank']} in administrative disbursements."
    )

    # Save file
    filename = f"{state.lower().replace(' ', '_')}_highway_report.docx"
    doc.save(filename)


for state in ["Florida", "Ohio"]:
    generate_report(state, state_data[state])
